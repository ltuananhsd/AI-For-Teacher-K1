"""RAG Ingestion Pipeline.

Flow for ALL document files:
  1. Extract raw text (per file type)
  2. Extract anchors from raw text (names/numbers that MUST survive)
  3. Clean noise (images, box-drawing, backslash escapes)
  4. LLM semantic chunking (temperature=0, strict entity-preservation rules)
  5. Coverage validation — compare chunks vs anchors → add raw fallback if miss
  6. Store all chunks to FAISS

Special handling (bypass LLM chunking, preserve scripted Q&A structure):
  - faq.md  → TSV parser (spreadsheet-export format: leading-tab rows)
  - FAQ xlsx → Excel row parser
"""
import os
import re
import logging
import json
import argparse
from typing import List

import pandas as pd
from docx import Document
from dotenv import load_dotenv
from pydantic import BaseModel, Field

from langchain_deepseek import ChatDeepSeek
from langchain_openai import OpenAIEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_core.documents import Document as LangchainDocument
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import JsonOutputParser
from langchain_text_splitters import RecursiveCharacterTextSplitter

from ingestion.text_cleaner import clean_text
from ingestion.anchor_extractor import extract_anchors
from ingestion.coverage_validator import validate_coverage, create_fallback_chunk
from ingestion.contextual_prepender import prepend_context

# --- Setup ---
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)
load_dotenv()

for _key in ["DEEPSEEK_API_KEY", "OPENAI_API_KEY"]:
    if not os.getenv(_key):
        logger.error(f"{_key} not found in .env")
        exit(1)

llm = ChatDeepSeek(model="deepseek-chat", temperature=0)
embeddings = OpenAIEmbeddings(model="text-embedding-3-small")
FAISS_PATH = "./faiss_index"

# --- LLM Output Schema ---
class SemanticChunk(BaseModel):
    title: str = Field(description="Short title describing the chunk content")
    content: str = Field(description="Full content, preserving all names/numbers verbatim from source text")
    summary: str = Field(description="Brief summary")
    tags: List[str] = Field(description="Relevant keywords")

class SemanticChunks(BaseModel):
    chunks: List[SemanticChunk]

# --- LLM Chunking Prompt (strict entity preservation) ---
_CHUNKING_SYSTEM = """You are a document analysis expert for a RAG (Retrieval-Augmented Generation) system.

TASK: Split the input text into semantic chunks for indexing into a vector database.

MANDATORY RULES — VIOLATION = WRONG:
1. VERBATIM COPY (CRITICAL): Person names, organizations, brands, statistics, prices, percentages, dates MUST be copied EXACTLY from the source — NO paraphrasing, abbreviating, or changing even 1 character
2. NO FABRICATION: Only use information that EXISTS in the source text. Absolutely no adding external information
3. NO INFORMATION LOSS: Every piece of information in the source MUST appear in at least 1 chunk — especially people, numbers, organization names
4. SELF-CONTAINED CHUNKS: Each chunk is a complete unit of information, readable independently and still understandable
5. GROUP BY TOPIC: Group paragraphs on the same topic into 1 chunk; do not split mechanically line by line
6. STRIP GARBAGE CHARACTERS: In chunk output content, remove meaningless characters such as: control characters, repeated dashes (---/===), meaningless special characters — but KEEP emojis/icons if present in source as they may carry visual meaning

{format_instructions}"""

# --- Raw Text Extractors ---
def _read_docx(file_path: str) -> str:
    try:
        doc = Document(file_path)
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    except Exception as e:
        logger.error(f"docx read error {file_path}: {e}")
        return ""

def _read_text_file(file_path: str) -> str:
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    except Exception as e:
        logger.error(f"text read error {file_path}: {e}")
        return ""

def _read_xlsx_as_text(file_path: str) -> str:
    """Convert xlsx sheets to plain text for LLM processing."""
    try:
        xls = pd.ExcelFile(file_path)
        parts = []
        for sheet in xls.sheet_names:
            df = pd.read_excel(file_path, sheet_name=sheet).fillna("")
            parts.append(f"[Sheet: {sheet}]")
            for _, row in df.iterrows():
                row_text = " | ".join(
                    f"{col}: {val}" for col, val in row.items() if str(val).strip()
                )
                if row_text:
                    parts.append(row_text)
        return "\n".join(parts)
    except Exception as e:
        logger.error(f"xlsx read error {file_path}: {e}")
        return ""

# --- LLM Semantic Chunker ---
def _llm_chunk(text: str, source_name: str) -> List[LangchainDocument]:
    """Call DeepSeek to split clean text into semantic chunks with entity preservation."""
    parser = JsonOutputParser(pydantic_object=SemanticChunks)
    prompt = ChatPromptTemplate.from_messages([
        ("system", _CHUNKING_SYSTEM),
        ("human", "Văn bản cần chia:\n\n{text}"),
    ])
    try:
        result = (prompt | llm | parser).invoke({
            "text": text,
            "format_instructions": parser.get_format_instructions(),
        })
        chunks_data = result.get("chunks", [])
        if not chunks_data:
            raise ValueError("LLM returned 0 chunks")

        docs = [
            LangchainDocument(
                page_content=c.get("content", ""),
                metadata={
                    "source": source_name,
                    "title": c.get("title", ""),
                    "summary": c.get("summary", ""),
                    "tags": ", ".join(c.get("tags", [])),
                    "type": "llm_semantic",
                },
            )
            for c in chunks_data
        ]
        logger.info(f"[{source_name}] LLM → {len(docs)} chunks")
        return docs
    except Exception as e:
        logger.error(f"[{source_name}] LLM chunking error: {e} — using fallback")
        return _fallback_split(text, source_name)

def process_md_semantic(file_path, text_content):
    """Use DeepSeek to extract clean text from Markdown and split into semantic chunks (with batching)."""
    logger.info(f"Processing markdown with DeepSeek: {os.path.basename(file_path)}")
    
    parser = JsonOutputParser(pydantic_object=SemanticChunks)
    
    prompt = ChatPromptTemplate.from_messages([
        ("system", "You are an expert tech writer and data extraction specialist. "
                   "The user will provide a section of a Markdown file. "
                   "Your task is to STRIP OUT all the code jargon and EXTRACT ONLY the human-readable text, knowledge, course syllabus, and explanations. "
                   "Split the extracted text into logical semantic sections. "
                   "For each section, provide a concise 'title', the pure text 'content' (no code tags), a brief 'summary', and relevant 'tags'. "
                   "Ensure ALL the actual course knowledge is preserved. "
                   "\n{format_instructions}"),
        ("human", "{text}")
    ])

    chain = prompt | llm | parser

    # Batching to avoid LLM output truncation on large files
    from langchain_text_splitters import RecursiveCharacterTextSplitter
    splitter = RecursiveCharacterTextSplitter(chunk_size=4000, chunk_overlap=200)
    raw_chunks = splitter.split_text(text_content)
    
    documents = []
    
    for i, raw_chunk in enumerate(raw_chunks):
        logger.info(f"Processing chunk {i+1}/{len(raw_chunks)} with DeepSeek...")
        try:
            result = chain.invoke({
                "text": raw_chunk,
                "format_instructions": parser.get_format_instructions()
            })
            
            chunks = result.get("chunks", [])
            for chunk in chunks:
                doc = LangchainDocument(
                    page_content=chunk.get('content', ''),
                    metadata={
                        "source": os.path.basename(file_path),
                        "title": chunk.get('title', 'Untitled'),
                        "summary": chunk.get('summary', ''),
                        "tags": ", ".join(chunk.get('tags', [])),
                        "type": "md_semantic"
                    }
                )
                documents.append(doc)
        except Exception as e:
            logger.error(f"Error processing semantic chunk {i+1}: {e}")

    if not documents:
        logger.warning(f"No chunks found in JSON response for {file_path}. Using fallback.")
        return _fallback_split(file_path, text_content)

    logger.info(f"Generated {len(documents)} clean semantic chunks for {os.path.basename(file_path)}")
    return documents

def _fallback_split(text: str, source_name: str) -> List[LangchainDocument]:
    # 512 chars ≈ 380 Vietnamese words — optimal for Vietnamese text retrieval
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=512, chunk_overlap=80,
        separators=["\n\n", "\n", ". ", ", ", " ", ""],
    )
    docs = splitter.create_documents([text])
    for doc in docs:
        doc.metadata = {"source": source_name, "type": "fallback"}
    logger.info(f"[{source_name}] Fallback split → {len(docs)} chunks")
    return docs


def deduplicate_new_file_chunks(
    new_file_docs: list,
    existing_faiss_path: str,
    embeddings_model,
    threshold: float = 0.92,
) -> tuple:
    """Compare only NEW file chunks against existing FAISS to detect duplicates.
    Old files' chunks are ALWAYS kept intact (no cross-file dedup).
    Returns (unique_docs, num_duplicates_removed)."""
    if not os.path.exists(existing_faiss_path) or not new_file_docs:
        return new_file_docs, 0
    try:
        old_vs = FAISS.load_local(
            existing_faiss_path, embeddings_model,
            allow_dangerous_deserialization=True
        )
    except Exception as e:
        logger.warning(f"Could not load existing FAISS for dedup: {e}")
        return new_file_docs, 0

    unique, dups = [], 0
    for doc in new_file_docs:
        results = old_vs.similarity_search_with_score(doc.page_content, k=1)
        if results:
            _, l2_dist = results[0]
            # Convert L2 distance to cosine similarity (valid for normalized OpenAI embeddings)
            cosine_sim = 1 - (l2_dist ** 2) / 2
            if cosine_sim > threshold:
                dups += 1
                logger.debug(f"Duplicate chunk removed (sim={cosine_sim:.3f}): {doc.page_content[:60]}")
                continue
        unique.append(doc)
    logger.info(f"Dedup: {dups} duplicates removed, {len(unique)} unique chunks kept")
    return unique, dups

def _split_md_sections(text: str) -> List[str]:
    """Split markdown by ## (H2) headers into focused sections for per-section LLM calls.
    Prevents LLM from skipping late sections in long documents.
    """
    # Split on lines that start with "##" (but not "###")
    parts = re.split(r'\n(?=## )', text)
    # Filter out tiny fragments (< 80 chars) that carry no real info
    return [p.strip() for p in parts if len(p.strip()) >= 80]

def _process_section(section_text: str, source_name: str, section_id: str) -> List[LangchainDocument]:
    """Process one section:
    clean → LLM chunk (children) → coverage validate → contextual prepend → return.

    Parent-Child strategy:
    - 1 parent chunk = full clean section (broad queries hit this directly)
    - N child chunks = LLM semantic splits (precise queries hit these)
    - Children carry parent_id so retriever can expand to parent at query time
    Fallback is scoped to ONLY this section (not the whole doc).
    """
    anchors = extract_anchors(section_text)  # from RAW before cleaning
    clean = clean_text(section_text)

    # --- Parent chunk: full clean section, no further splitting ---
    section_header = clean.splitlines()[0][:80] if clean else source_name
    parent_chunk = LangchainDocument(
        page_content=clean,
        metadata={
            "source": source_name,
            "title": f"[Section] {section_header}",
            "type": "section_parent",
            "chunk_level": "parent",
            "section_id": section_id,
        },
    )

    # --- Child chunks: LLM semantic splits ---
    children = _llm_chunk(clean, source_name)

    # Coverage check — add focused fallback if any anchor is missing
    missing = validate_coverage(children, anchors, source_name)
    if missing:
        children.append(create_fallback_chunk(missing, section_text, source_name))

    # Tag children with parent_id so main.py can expand to parent on retrieval
    for child in children:
        child.metadata["chunk_level"] = "child"
        child.metadata["parent_id"] = section_id

    # Contextual prepending on children only (parent is already full section)
    children = prepend_context(children, section_text, llm, source_name)

    # Return parent first, then children — both go into FAISS
    return [parent_chunk] + children

# --- Universal Document Pipeline ---
def process_document(file_path: str) -> List[LangchainDocument]:
    """
    Universal pipeline for any text document:
    raw text → (section split for .md) → per-section LLM chunk + coverage validate → FAISS

    Markdown files are split by ## headers first so LLM never processes >1 section
    at a time — prevents it from dropping later sections (e.g. giảng viên info).
    """
    source_name = os.path.basename(file_path)
    ext = os.path.splitext(file_path)[1].lower()

    # Step 1: Extract raw text
    if ext == ".docx":
        raw = _read_docx(file_path)
    elif ext in [".md", ".txt"]:
        raw = _read_text_file(file_path)
    elif ext in [".xlsx", ".xls"]:
        raw = _read_xlsx_as_text(file_path)
    else:
        logger.warning(f"Unsupported format: {ext} — skipping {source_name}")
        return []

    if not raw.strip():
        logger.warning(f"[{source_name}] Empty content — skipping")
        return []

    # Step 2: For markdown — split by sections, process each independently
    if ext in [".md", ".txt"]:
        sections = _split_md_sections(raw)
        logger.info(f"[{source_name}] Split into {len(sections)} sections")
        all_chunks: List[LangchainDocument] = []
        for i, section in enumerate(sections):
            header_preview = section.splitlines()[0][:60] if section else ""
            logger.info(f"[{source_name}] Section {i+1}/{len(sections)}: {header_preview}")
            section_id = f"{source_name}::section_{i}"
            all_chunks.extend(_process_section(section, source_name, section_id))
        return all_chunks

    # Step 3: For other file types — process whole document at once
    anchors = extract_anchors(raw)
    logger.info(
        f"[{source_name}] Anchors: {len(anchors['names'])} names, "
        f"{len(anchors['numbers'])} numbers"
    )
    clean = clean_text(raw)
    chunks = _llm_chunk(clean, source_name)
    missing = validate_coverage(chunks, anchors, source_name)
    if missing:
        chunks.append(create_fallback_chunk(missing, raw, source_name))
    return chunks

# --- Special FAQ Parsers (scripted Q&A — bypass LLM chunking) ---
def _process_faq_md(file_path: str) -> List[LangchainDocument]:
    """
    Parse FAQ markdown exported from spreadsheet (TSV format with leading tab).
    Each data line: \tQuestion\tAnswer → parts = ['', question, answer]
    """
    source_name = os.path.basename(file_path)
    docs = []
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            lines = f.readlines()
        for index, line in enumerate(lines):
            parts = line.split('\t')
            if len(parts) < 3:
                continue
            q = parts[1].strip()
            a = parts[2].strip()
            if not q or not a or "Câu Hỏi" in q or q.lower() == 'nan' or len(q) < 5 or len(a) < 5:
                continue
            if len(q) > 2 and q[0].isdigit() and ". " in q[:5]:
                q = q.split(". ", 1)[1]
            docs.append(LangchainDocument(
                page_content=f"QUESTION: {q}\nANSWER: {a}",
                metadata={"source": source_name, "type": "script_faq", "row": index},
            ))
        logger.info(f"[{source_name}] {len(docs)} FAQ pairs loaded")
    except Exception as e:
        logger.error(f"FAQ md error {file_path}: {e}")
    return docs

def _process_faq_xlsx(file_path: str) -> List[LangchainDocument]:
    """Parse FAQ Excel: Q in col 0, A in col 1, skip header rows."""
    source_name = os.path.basename(file_path)
    docs = []
    try:
        df = pd.read_excel(file_path, header=None)
        for index, row in df.iterrows():
            if index < 2 or len(row) < 2:
                continue
            q = str(row[0]).strip() if pd.notna(row[0]) else ""
            a = str(row[1]).strip() if pd.notna(row[1]) else ""
            if not q or not a or q.lower() == 'nan' or a.lower() == 'nan':
                continue
            if q[0].isdigit() and ". " in q:
                q = q.split(". ", 1)[1]
            docs.append(LangchainDocument(
                page_content=f"QUESTION: {q}\nANSWER: {a}",
                metadata={"source": source_name, "type": "script_faq", "row": index},
            ))
        logger.info(f"[{source_name}] {len(docs)} FAQ pairs loaded")
    except Exception as e:
        logger.error(f"FAQ xlsx error {file_path}: {e}")
    return docs

# --- Main ---
def main():
    parser = argparse.ArgumentParser(description="RAG Ingestion Pipeline")
    parser.add_argument("--new-file", default=None,
                        help="Filename of the newly uploaded file (for dedup scope)")
    args = parser.parse_args()

    target_dirs = ["tailieuduan"]
    all_documents: List[LangchainDocument] = []
    report = []

    for target_dir in target_dirs:
        if not os.path.exists(target_dir):
            logger.warning(f"Directory not found: {target_dir}")
            continue
        logger.info(f"Scanning: {target_dir}")

        files = sorted(
            f for f in os.listdir(target_dir)
            if os.path.isfile(os.path.join(target_dir, f))
        )
        for file_name in files:
            file_path = os.path.join(target_dir, file_name)
            name_lower = file_name.lower()
            ext = os.path.splitext(file_name)[1].lower()

            # Skip binary/image files
            if ext in [".jpg", ".jpeg", ".png", ".gif", ".svg", ".webp", ".pdf"]:
                continue

            # FAQ files → scripted Q&A parser (exact answers, high priority in retrieval)
            if "faq" in name_lower and ext == ".md":
                docs = _process_faq_md(file_path)
            elif "faq" in name_lower and ext in [".xlsx", ".xls"]:
                docs = _process_faq_xlsx(file_path)
            # All other text documents → universal LLM pipeline
            elif ext in [".docx", ".md", ".txt", ".xlsx", ".xls"]:
                docs = process_document(file_path)
            else:
                logger.warning(f"Skipping unsupported: {file_name}")
                continue

            report.append({
                "file": file_name,
                "chunks": len(docs),
                "types": list(set(d.metadata.get("type", "?") for d in docs)),
            })
            all_documents.extend(docs)

    if not all_documents:
        logger.warning("No documents processed.")
        print(f'DEDUP_STATS:{json.dumps({"dedup_stats": {"total_before_dedup": 0, "duplicates_removed": 0, "unique_chunks_added": 0}})}')
        return

    # --- Deduplication (only for newly uploaded file's chunks) ---
    dups_removed = 0
    total_before = 0
    if args.new_file:
        new_file_docs = [
            d for d in all_documents
            if os.path.basename(d.metadata.get("source", "")) == args.new_file
        ]
        existing_docs = [
            d for d in all_documents
            if os.path.basename(d.metadata.get("source", "")) != args.new_file
        ]
        total_before = len(new_file_docs)
        filtered_new, dups_removed = deduplicate_new_file_chunks(
            new_file_docs, FAISS_PATH, embeddings
        )
        all_documents = existing_docs + filtered_new
        logger.info(
            f"Dedup summary: {dups_removed}/{total_before} new chunks were duplicates. "
            f"Final index size: {len(all_documents)} chunks."
        )

    # Ingest to FAISS
    logger.info(f"Indexing {len(all_documents)} chunks to FAISS...")
    try:
        vectorstore = FAISS.from_documents(all_documents, embedding=embeddings)
        vectorstore.save_local(FAISS_PATH)
        logger.info("Ingestion complete → ./faiss_index")
    except Exception as e:
        logger.error(f"FAISS indexing error: {e}")
        return

    # Ingestion report
    print("\n=== INGESTION REPORT ===")
    for r in report:
        types_str = ", ".join(r["types"])
        print(f"  {r['file']:<35} {r['chunks']:>3} chunks  [{types_str}]")
    print(f"  {'TOTAL':<35} {len(all_documents):>3} chunks indexed\n")

    # Dedup stats (parsed by rag_api.py)
    dedup_stats = {
        "total_before_dedup": total_before,
        "duplicates_removed": dups_removed,
        "unique_chunks_added": len(all_documents),
    }
    print(f'DEDUP_STATS:{json.dumps({"dedup_stats": dedup_stats})}')

if __name__ == "__main__":
    main()
